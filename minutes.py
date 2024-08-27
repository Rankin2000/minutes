from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

os.system("clear")
os.chdir("/home/stuart/Documents/LaptopFiles/Minutes/minute/")
#Old Project or New
project = False
os.chdir("Minutes")
while project == False:
    project = input("New Project? (y/n)")
    if project == "y":
        project = True
    else:
        project = False

    if project:
        projectName = input("Please enter the name of the new project? ")
        projectName = projectName.strip()
        projectName = projectName.title()
        projectName = projectName.replace(" ", "")
        os.mkdir(projectName)
    else:
        directory = os.listdir()
        print("\nWhich project do you want to use?")
        if len(directory):
            i = 1
            for x in directory:
                print(str(i)+". "+x)
                i = i + 1
            temp = input("\nPlease enter the number for your project ")
            project = True
            projectName = directory[int(temp)-1]
        else:
            print("There are no projects to choose from.")
os.chdir(projectName)
print(os.listdir())
document = Document()

title = projectName

date = input("Is the meeting today? (y/n) ")
if date == "y":
    date = datetime.now().date()
    #date = str(date)
else:
    date = input("Please enter the date the meeting was on? (yyyy-mm-dd) ")
#date = date.replace("-", "/")
time = input("What time did the meeting start? (24hr xx:yy) ")

location = input("Where did the meeting take place? ")

present = [] 
while True:
    person = input("Enter a name or 'done' ")
    if person == "done":
        break
    present.append(person)

absent = []

while True:
    person = input("Was there any absences? Enter name or 'done' ")
    if person == "done":
        break
    absent.append(person)


#quorum = input("Is/was a quorum present? (y/n) ")
if len(present) > len(absent):
    quorum = True
else:
    quorum = False




items = []
notes = []
atmos = []
nextitems = []
nextmeeting = ""
while True:
    
    if items:
        print("1. Agenda\n")
    else:
        print("1. Agenda - Not Done\n") 
    if notes:
        print("2. Notes\n")
    else:
        print("2. Notes - Not Done\n")
    if atmos:
        print("3. Atmosphere and Challenges\n")
    else:
        print("3. Atmosphere and Challenges - Not Done\n")
    if nextmeeting != "":
        print("4. Next Meeting\n")
    else:
        print("4. Next Meeting - Not Done\n")
    print("5. Finish and Save\n") 

    temp = input()
    if temp == "1":
        print("Enter an item for the agenda\nEnter 'done' when finished ")
        while True:
            item = input()
            if item != "":
                if item == "done":
                    break
                items.append(item)
    elif temp == "2":
        print("Enter a note, enter 'done' when finished note")
        while True:
            note = input()
            if note  == "done":
                break
            notes.append(note)
    elif temp == "3":
        print("Atmosphere and Challenges")
        print("Enter 'done'  to finish")
        while True:
            atmosphere = input()
            if atmosphere == "done":
                break
            atmos.append(atmosphere)
    elif temp == "4":
        print("Next Meeting details")
        nextmeeting = input("Is there meeting planned? (y/n) ")
        if nextmeeting == "y":
            week = input("Is the meeting the same time next week? (y/n)")
            if week == "y":
                date = datetime.now().date()
                today = date.strftime("%d")
                currentmonth = date.strftime("%m")
                year = int(date.strftime("%Y"))
                nextday = int(today) + 7
                nextmonth = currentmonth
                if nextday ==  28 or nextday >= 30:
                    if currentmonth == "02":
                        if ((year%400 == 0) or ((year%4 == 0) and (year%100 != 0))):
                            nextday -= 29
                        else:
                            nextday -= 28
                    if currentmonth == "04" or currentmonth == "06" or currentmonth == "09" or currentmonth == "11":
                        nextday -= 30
                    else:
                        if currentmonth == "12":
                            year += 1
                            nextmonth = "01"
                        if nextday > 31:
                            nextday -= 31
                if nextday < 10:
                    nextday = 0 + str(nextday)
                nextdate = str(year)+"-"+str(nextmonth)+"-"+str(nextday)

        else:
            nextdate = input("What day is the meeting? (yyyy-mm-dd)")
        
        nexttime = input("What time is it starting? (xx:yy)")
        nextlocation = input("Where is the location? ")
        print("Enter an item\nEnter 'done' when finished ")
        while True:
            nextitem = input()
            if nextitem != "":
                if nextitem == "done":
                    break
                nextitems.append(nextitem)
    elif temp == "5":
        if items and notes and atmos and nextmeeting != "":
            finish = input("Are you ready to finish? (y/n)")
            if finish == "y":
                break
        else:
            print("Please finish all sections")

items.reverse()
notes.reverse()
atmos.reverse()
nextitems.reverse()

if date == datetime.now().date():
    endtime = input("If the meeting is finished enter the time (xx:yy) or blank to return ")
else:
    endtime = input("When did the meeting finish? (xx:yy) ")


font = document.styles["Normal"].font
font.name = "Arial"
font.size = Pt(14)
document.add_heading(title, level=0)
info = document.add_paragraph(str(date)+"\nStart Time: "+str(time)+"\nEnd Time: "+str(endtime)+"\n"+location)
info.alignment = WD_ALIGN_PARAGRAPH.RIGHT

attendees = "Attendees: "
present.reverse()

while present:
    attendees += present.pop()
    attendees += ", "
attendees = attendees[:-2]

absentees = "Absentees: "
absent.reverse()
while absent:
    absentees += absent.pop()
    absentees += ", "
absentees = absentees[:-2]

if quorum:
    info2 = document.add_paragraph(attendees+"\n"+absentees+"\nQuorum Present")
else:
    info2 = document.add_paragraph(attendees+"\n"+absentees+"\nQuorum Not Present")

#agenda = ""
document.add_heading("Agenda", level=1)
while items:
    document.add_paragraph(items.pop(), style="List Bullet")

document.add_heading("Notes", level=1)
noted = ""
while notes:
    noted += notes.pop()
    noted += "\n"
document.add_paragraph(noted)

document.add_heading("Atmosphere and Challenges", level=1)
atmosphere = ""
while atmos:
    atmosphere += atmos.pop()
    atmosphere += "\n"
document.add_paragraph(atmosphere)

document.add_heading("Next Meeting", level=1)
if nextmeeting != "y":
    document.add_paragraph("There is no meeting planned.")
else:
    meeting = document.add_paragraph(nextdate+"\nStart Time: "+str(time)+"\n"+nextlocation)
    meeting.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_heading("Next Meeting Agenda", level=1)
    while nextitems:
        document.add_paragraph(nextitems.pop(), style="List Bullet")
document.save(str(date)+time.replace(":","")+".docx")
